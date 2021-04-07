from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Cm, Inches, Pt

from .. import reports as rp
from .. import styles as st
from . import fileops as fl
from .__init__ import *

log = getlog(__name__)


class WordReport():
    def __init__(self, **kw) -> None:
        doc = Document()
        font_name = 'Calibri'
        font_size = Pt(10)
        ext = 'docx'

        font_props = dict(
            size=font_size,
            name=font_name)
        
        tables = {}

        f.set_self(vars())

        # set default paragraph font styles
        self.set_style_props('Normal', font_props)
        self.set_style_props('No Spacing', font_props)
        self.set_style_props('Normal Table', font_props)

        # small font table
        s = doc.styles.add_style('SmallFont', WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = doc.styles['No Spacing']
        self.set_style_props('SmallFont', dict(size=Pt(6)))

        self.doc.styles['Normal'].paragraph_format.line_spacing = 1
        self.doc.styles['Heading 1'].paragraph_format.space_before = Pt(6)
        
        self.set_margins()
    
    def set_header_footer(self):
        
        p = f.resources / 'reports/images'
        p_head = p / 'SMS Logo.png'
        p_foot = p / 'SMS Footer.png'

        section = self.doc.sections[0]
        header = section.header.paragraphs[0]
        footer = section.footer.paragraphs[0]

        header.paragraph_format.space_after = Pt(12)

        header.add_run().add_picture(str(p_head), height=Inches(0.4))
        footer.add_run().add_picture(str(p_foot))
    
    def set_style_props(self, name, props):
        style = self.doc.styles[name]
        font = style.font

        for k, v in props.items():
            setattr(font, k, v)
    
    def set_margins(self):
        """Set full document margins"""
        doc = self.doc
        sections = doc.sections

        for section in sections:
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

    def add_df(self, df, name: str=None, header: bool=True, index: bool=False, m_bg: dict=None, m_text: dict=None, header_color: str=f.config['color']['thead'], m_hdrs: dict=None, **kw):
        """Add table to word doc from df

        Parameters
        ----------
        df : pd.DataFrame
            dataframe to add
        name : str, optional
            name to save table as, by default 'table_1'
        """

        if name is None:
            name = f'table_{len(self.tables.keys()) + 1}'
        
        hdr = 1 if header else 0
        idx = 1 if index else 0

        num_cols = df.select_dtypes('number').columns.tolist()
        
        tbl = self.doc.add_table(
            rows=df.shape[0] + hdr,
            cols=df.shape[1] + idx)
        
        # save table
        self.tables[name] = tbl

        # add column labels
        if header:
            for i, col_name in enumerate(df.columns):
                cell = tbl.rows[0].cells[i + idx]
                cell.text = col_name
                self.set_cell_color(cell, bg=header_color)
                self.set_cell_font(cell=cell, align=WD_ALIGN_PARAGRAPH.CENTER)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        if index:
            # set index name
            if header:
                cell = tbl.rows[0].cells[0]
                cell.text = df.index.name
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            for i, row_name in enumerate(df.index):
                cell = tbl.rows[i + hdr].cells[0]
                cell.text = row_name
        
        # add data
        for i, row in enumerate(df.itertuples(index=False)):
            row_name = df.index[i]

            for j, val in enumerate(row):
                cell = tbl.rows[i + hdr].cells[j + idx]
                col_name = df.columns[j]

                int_col = pd.api.types.is_integer_dtype(df.dtypes[col_name])

                # set number formats, don't add None
                if not val is None:
                    if int_col:
                        val = f'{val:,.0f}'

                    cell.text = str(val)

                # align numeric cols right
                if col_name in num_cols:
                    self.set_cell_font(cell=cell, align=WD_ALIGN_PARAGRAPH.RIGHT)

                if not m_bg is None:

                    # convert btwn original/display column names
                    if not m_hdrs is None:
                        col_name = m_hdrs.get(col_name, col_name)

                    color = m_bg.get(col_name, {}).get(row_name, None)
                    if color:
                        self.set_cell_color(cell, bg=color)

        self.set_table_cell_style(tbl=tbl, **kw)

        return tbl
    
    def set_cell_color(self, cell, bg=None, text=None):
        """Set table cell bg and text color"""
        cell_pr = cell._element.tcPr
        cl_shading = OxmlElement('w:shd')

        if bg:
            cl_shading.set(qn('w:fill'), bg)
        
        if text:
            cl_shading.set(qn('w:color'), text)

        cell_pr.append(cl_shading)

    def set_table_cell_style(self, tbl, style_name='No Spacing'):
        """Set font size/name for full table from Paragraph style"""
        for row in tbl.rows:
            for cell in row.cells:
                p = cell.paragraphs[0]
                p.style = self.doc.styles[style_name]
    
    def autofit_table_window(self, tbl):
        """
        <w:tbl>
            <w:tblPr>
                <w:tblW w:type="pct" w:w="5000"/>
        """
        width = OxmlElement('w:tblW')
        width.set(qn('w:type'), 'pct')
        width.set(qn('w:w'), '5000')
        tbl._tbl.tblPr.append(width)

    def set_table_margins(self, tbl, width: dict=None):
        """<w:tbl>
			<w:tblPr>
				<w:tblStyle w:val="LightShading"/>
				<w:tblW w:type="auto" w:w="0"/>
				<w:tblCellMar>
					<w:left w:type="dxa" w:w="63"/>
					<w:right w:type="dxa" w:w="63"/>
				</w:tblCellMar>
				<w:tblLook w:firstColumn="1" w:firstRow="1" w:lastColumn="0" w:lastRow="0" w:noHBand="0" w:noVBand="1" w:val="04A0"/>
			</w:tblPr>"""
        
        # 67 = Cm(0.11) ...?
        if width is None:
            width = dict(left=67, right=67)

        margins = OxmlElement('w:tblCellMar')

        for side, w in width.items():
            margin = OxmlElement(f'w:{side}')
            margin.set(qn('w:w'), str(w))
            margin.set(qn('w:type'), 'dxa')

            margins.append(margin)

        tbl._tbl.tblPr.append(margins)
    
    def add_tbl_border(self, tbl):
        """Add table bottom border with OxmlElement"""
        borders = OxmlElement('w:tblBorders')
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), '4')
        borders.append(bottom_border)

        tbl._tbl.tblPr.append(borders)
    
    def set_cell_font(self, cell, props: dict=None, align=None):
        """Set cell font/properties"""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)

        if not align is None:
            p.alignment = align

        if not props is None:
            for run in p.runs:
                font = run.font
                for name, val in props.items():
                    setattr(font, name, val)
    
    def bold_column(self, tbl, cols: list):
        """Set table columns to bold"""
        for row in tbl.rows:
            for i in cols:
                cell = row.cells[i]
                self.set_cell_font(cell=cell, props=dict(bold=True))
    
    def bold_header(self, text):
        doc = self.doc

        p = doc.add_paragraph()
        run = p.add_run(f'{text.title()}:')
        run.bold = True
        run.add_break()
    
    def add_pictures(self, pics: list):
        doc = self.doc
        doc.add_page_break()
        doc.add_heading('Pictures', level=2)

        for i, pic in enumerate(pics):
            
            doc.add_picture(str(pic), height=Inches(4))

            # center picture
            p = doc.paragraphs[-1]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add figure caption
            cap = doc.add_paragraph(f'Figure ', style='Caption')
            self.add_fig_caption(cap)
            cap.add_run(' - ')
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_fig_caption(self, paragraph):
        """Add figure caption to image with auto updating numbers
        - User must select all (cmd/ctrl + A), then F9 to update fig captions"""
        run = paragraph.add_run()
        r = run._r
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar)
        instrText = OxmlElement('w:instrText')
        instrText.text = ' SEQ Figure \* ARABIC'
        r.append(instrText)
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        r.append(fldChar)

    def save(self, p: Path=None):

        if p is None:
            p = self.p

        self.doc.save(p)
    
    def show(self):
        fl.open_folder(self.p_rep)

    def create_word(self, p_base=None, check_overwrite=False, **kw):
        if p_base is None:
            p_base = Path.home() / 'Desktop'
        
        p = p_base / f'{self.title}.docx'

        self.doc.save(str(p))
        self.p_rep = p
        return self   



class FailureReportWord(WordReport, rp.FailureReport):
    """Failure report based on both rp.FailureReport (for data/initialization) and Word report"""
    def __init__(self, query_oil=None, **kw):
        super().__init__(**kw)
        rp.FailureReport.__init__(self, **kw)

        p = f.desktop / 'component_failure.docx'

        f.set_self(vars())

    def create_word(self, check_overwrite=True, **kw):
        doc = self.doc

        self.set_header_footer()

        p = doc.add_heading('Component Failure Summary', level=1)
        p.runs[0].font.size = Pt(20)
        p.paragraph_format.space_after = Pt(12)

        tbl = self.add_df(df=self.df_head, name='header', header=False)
        self.bold_column(tbl=tbl, cols=(0, 2))
        self.add_tbl_border(tbl)

        self.add_body()

        if not self.query_oil is None:
            self.add_oil_samples()

        if self.pictures:
            self.add_pictures(pics=sorted(self.pictures))

        return super().create_word(
            p_base=self.ef._p_event,
            check_overwrite=check_overwrite,
            **kw)

    def add_oil_samples(self):
        """Add oil samples as two split tables for more room"""
        query = self.query_oil

        df = query.df \
            .assign(
                sample_date=lambda x: x.sample_date.astype(str),
                oil_changed=lambda x: x.oil_changed.astype(str)) \
            .rename(columns=dict(sample_date='Sample_Date')) \
            .set_index('Sample_Date')

        if 'opc' in df.columns:
            df['opc'] = df.opc.str.replace(' ', '')

        style = df.style \
            .pipe(st.default_style, outlook=True) \
            .pipe(query.update_style)

        # set vals for title
        unit = df.unit.iloc[0]
        component = df.component_id.iloc[0]
        modifier = df.modifier.iloc[0]

        flag_cols = [col for col in df.columns if '_fg' in col]
        drop_cols = ['unit', 'component_id', 'modifier'] #, 'unit_smr']
        # , 'sliding', 'fatigue', 'non_metal', 'fibers', 'cutting_mean', 'sliding_mean', 'fatigue_mean', 'non_metal_mean']
        drop_cols.extend(flag_cols)

        # convert headers for snake vs title cols
        cols = df.columns.tolist()
        cols_nice = f.lower_cols(cols, title=True)
        m_hdrs = {nice_col: col for col, nice_col in zip(cols, cols_nice)}

        df = df \
            .drop(columns=drop_cols) \
            .pipe(f.lower_cols, title=True)

        m_bg, m_text = f.convert_stylemap_index_color(style, as_qt=False, first_val=False)

        split = (df.shape[1] // 2) + 8

        m_tbls = dict(
            oil1=df.iloc[:, :split],
            oil2=df.iloc[:, split:])
        
        # add header
        header = f'{unit} - {component}'
        if not modifier is None:
            header = f'{header}, {modifier}'

        self.doc.add_heading('Oil Samples', level=2)
        self.doc.add_paragraph().add_run(header.title())

        for name, df in m_tbls.items():
            tbl = self.add_df(df, name=name, m_bg=m_bg, m_hdrs=m_hdrs, style_name='SmallFont', index=True)
            self.set_table_margins(tbl=tbl, width=dict(left=50, right=50))
            self.bold_column(tbl=tbl, cols=(0,))
            self.doc.add_paragraph()
    
    def add_body(self):
        doc = self.doc
        doc.add_paragraph()
        
        for header, text in self.body.items():
            self.bold_header(header)
            p = doc.paragraphs[-1]
            p.add_run(text)


# from smseventlog.utils import word as wd

# query = qr.OilSamplesReport(unit='F302', component='WHEEL MOTOR', modifier='LEFT')
# df = query.get_df()
# style = df.style.pipe(query.update_style)

# uid = 161349647756
# rep = wd.FailureReportWord.example(uid=uid, style_oil=style)
# rep.create_word()
# rep.show()
